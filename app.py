from flask import Flask, render_template, request, send_file
import os
import cv2
import numpy as np
from google.cloud import vision
from docx import Document
from docx.shared import Pt, Inches
import io

# Initialize Flask app
app = Flask(__name__)

# Set the upload folder
UPLOAD_FOLDER = "uploads"
RESULT_FOLDER = "results"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULT_FOLDER, exist_ok=True)

# Path to Tesseract executable
# Update or comment out if not required
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Function to process the uploaded image
def process_image(image_path, output_path):
    client = vision.ImageAnnotatorClient()

    # Read image
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    # Text detection
    response = client.document_text_detection(image=image)

    # Extract text blocks
    blocks = []
    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            block_text = ''.join([
                symbol.text for paragraph in block.paragraphs
                for word in paragraph.words for symbol in word.symbols
            ])
            vertices = [(vertex.x, vertex.y) for vertex in block.bounding_box.vertices]
            blocks.append({'text': block_text, 'vertices': vertices})

    # Sort blocks by position
    sorted_blocks = sorted(blocks, key=lambda b: (b['vertices'][0][1], b['vertices'][0][0]))

    # Create Word document
    doc = Document()
    for block in sorted_blocks:
        paragraph = doc.add_paragraph(block['text'])
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(block['vertices'][0][1] / 10)  # Approximate vertical spacing

    doc.save(output_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return "No file uploaded", 400

    file = request.files['file']
    if file.filename == '':
        return "No file selected", 400

    # Save the uploaded file
    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(input_path)

    # Process the file
    output_path = os.path.join(RESULT_FOLDER, "output.docx")
    process_image(input_path, output_path)

    return send_file(output_path, as_attachment=True, download_name="output.docx")

if __name__ == '__main__':
    app.run(debug=True)
