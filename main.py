import cv2
import numpy as np
import pytesseract
from PIL import Image
from google.cloud import vision
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os

# Path to Tesseract executable (if not in system PATH)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Step 1: Extract text and layout with Google Vision API
def extract_text_with_layout(image_path):
    client = vision.ImageAnnotatorClient()
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)
    response = client.document_text_detection(image=image)
    blocks = []
    for page in response.full_text_annotation.pages:
        for block in page.blocks:
            block_text = ''.join([
                symbol.text for paragraph in block.paragraphs
                for word in paragraph.words for symbol in word.symbols
            ])
            vertices = [(vertex.x, vertex.y) for vertex in block.bounding_box.vertices]
            blocks.append({'text': block_text, 'vertices': vertices})
    return blocks

# Step 2: Detect logos using Google Vision API with background transparency
def extract_logos(image_path, output_dir):
    client = vision.ImageAnnotatorClient()
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)
    response = client.logo_detection(image=image)

    logos = []
    original_image = cv2.imread(image_path)

    for i, annotation in enumerate(response.logo_annotations):
        logo_description = annotation.description
        vertices = [(vertex.x, vertex.y) for vertex in annotation.bounding_poly.vertices]

        # Save the detected logo with transparent background
        if len(vertices) == 4:  # Ensure valid bounding box
            x_min = min(v[0] for v in vertices)
            y_min = min(v[1] for v in vertices)
            x_max = max(v[0] for v in vertices)
            y_max = max(v[1] for v in vertices)
            cropped_logo = original_image[y_min:y_max, x_min:x_max]

            # Convert to RGBA format
            cropped_logo_rgba = cv2.cvtColor(cropped_logo, cv2.COLOR_BGR2BGRA)
            white_mask = (cropped_logo[:, :, 0] > 200) & (cropped_logo[:, :, 1] > 200) & (cropped_logo[:, :, 2] > 200)
            cropped_logo_rgba[white_mask] = [0, 0, 0, 0]  # Set white background to transparent

            # Save the logo
            output_file = os.path.join(output_dir, f"logo_{i}.png")
            cv2.imwrite(output_file, cropped_logo_rgba)
            logos.append({'description': logo_description, 'path': output_file, 'position': (x_min, y_min)})

    return logos

# Step 3: Sort blocks for alignment
def sort_blocks(blocks):
    # Sort blocks by top-left y coordinate, then by x coordinate
    return sorted(blocks, key=lambda b: (b['vertices'][0][1], b['vertices'][0][0]))

# Step 4: Check for borders
def has_border(block, image):
    vertices = block['vertices']
    x_min = min(v[0] for v in vertices)
    y_min = min(v[1] for v in vertices)
    x_max = max(v[0] for v in vertices)
    y_max = max(v[1] for v in vertices)

    # Extract the region of interest (ROI)
    roi = image[y_min:y_max, x_min:x_max]

    # Check for border by analyzing the edges
    edges = cv2.Canny(roi, 100, 200)
    return np.sum(edges) > 0

# Step 5: Create a Word document with aligned content and logos
def create_word_document(blocks, logos, output_path, image_path):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)

    sorted_blocks = sort_blocks(blocks)
    image = cv2.imread(image_path)

    for block in sorted_blocks:
        paragraph = doc.add_paragraph(block['text'])
        font = paragraph.runs[0].font
        font.size = Pt(10)  # Adjust font size as needed

        if has_border(block, image):
            # Add borders if the block has borders
            tc_pr = paragraph._element.get_or_add_pPr()
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Border size
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Black color
                tc_pr.append(border)

    # Add logos to the document
    for logo in logos:
        doc.add_paragraph(f"Logo: {logo['description']}")
        doc.add_picture(logo['path'], width=Inches(1.5))  # Adjust size as needed

    doc.save(output_path)
    print(f"Document saved to {output_path}")

# Main execution
if __name__ == "__main__":
    # Path to your image file
    image_path = 'images/sample_image4.jpg.jpeg'  # Adjust path accordingly
    output_dir = 'extracted_elements'  # Directory to save extracted logos/signatures
    os.makedirs(output_dir, exist_ok=True)

    # Extract text and layout
    blocks = extract_text_with_layout(image_path)
    print("Extracted Text and Layout:")
    for block in blocks:
        print(block)

    # Extract logos
    logos = extract_logos(image_path, output_dir)
    print("Extracted Logos:")
    for logo in logos:
        print(logo)

    # Create combined document
    output_path = "output_combined.docx"
    create_word_document(blocks, logos, output_path, image_path)
    print(f"Document saved as {output_path}")
